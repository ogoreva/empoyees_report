import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from docx.table import _Cell
from loguru import logger

# CONSTANTS
BG_GRAY_50_COLOR = '4d4d4d'
BG_GRAY_25_COLOR = '999999'
F_BLACK_COLOR = (0, 0, 0)
F_WHITE_COLOR = (255, 255, 255)


def set_text(
        cell: _Cell,
        text: str,
        bg_color: str | None = None,
        f_color: tuple[int, int, int] = F_BLACK_COLOR,
        bold: bool | None = None,
        alignment: WD_ALIGN_PARAGRAPH | None = None
) -> None:
    """
    Set background color to specified color. Color should be in hexadecimal format without leading sharp(#)
    :param cell:
    :param text:
    :param bg_color: Background color in hex format. Example: 4d4d4d
    :param f_color: Font color in rgb format. Example (16,16,16)
    :return:
    """
    # get paragraphs
    p = cell.paragraphs[0]
    if bg_color is not None:
        # set bg-color if specified
        # Unfortunately, python-docx package does not have public API
        # for setting background color. It requires actions with pure XML tags which you may see below
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        p.paragraph_format.element.get_or_add_pPr()
        p.paragraph_format.element.pPr.append(shd)
    # Specify Font color as an RGB object
    p.add_run(text).font.color.rgb = RGBColor(*f_color)
    # Make font is bold
    if bold is not None:
        p.runs[0].font.bold = bold
    # Align text if it needed
    if alignment is not None:
        p.alignment = alignment


def docx_exporter(items: list[dict], save_to: pathlib.Path):
    """
    Create Docx objects and export data into it
    :param items:
    :return:
    """
    logger.debug("Exporter executed")
    # Create document
    doc = Document()
    # specify styles for doc
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    # add leading title
    head = doc.add_paragraph('Отчет по загрузке')
    doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Create table
    table = doc.add_table(1, 2, 'TableGrid')
    head_cells = table.rows[0].cells

    # Fill table heads
    for i, item in enumerate(['Отдел', 'Количество задач']):
        set_text(head_cells[i], item, BG_GRAY_50_COLOR, F_WHITE_COLOR, True, WD_ALIGN_PARAGRAPH.CENTER)

    # Fill table with data
    for row in items:
        cells = table.add_row().cells
        set_text(cells[0], f"Отдел {row['department']}", BG_GRAY_25_COLOR, F_BLACK_COLOR, True)
        set_text(cells[1], str(row['total_task']), BG_GRAY_25_COLOR, F_BLACK_COLOR, True, WD_ALIGN_PARAGRAPH.CENTER)

        for i in sorted(row['employees'], key=lambda x: x["tasks"], reverse=True):
            cells = table.add_row().cells
            set_text(cells[0], str(i['name']))
            set_text(cells[1], str(i['tasks']), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    logger.debug(f"Save report to {save_to.absolute()}")
    # Save document
    doc.save(str(save_to))
